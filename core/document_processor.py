# Copyright (c) 2024-2026 JARVIS Project
# Licensed under the Apache License, Version 2.0 (the "License");
# you may not use this file except in compliance with the License.
# You may obtain a copy of the License at
#
#     http://www.apache.org/licenses/LICENSE-2.0
#
# Unless required by applicable law or agreed to in writing, software
# distributed under the License is distributed on an "AS IS" BASIS,
# WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
# See the License for the specific language governing permissions and
# limitations under the License.
from __future__ import annotations

import asyncio
import base64
import logging
import threading
from dataclasses import dataclass
from datetime import datetime, timedelta
from pathlib import Path
from typing import Any

logger = logging.getLogger(__name__)

TOKEN_LIMIT = 8_000
MAX_CHUNKS = 10
CHUNK_CHARS = 4_000
MAX_UPLOAD_MB = 50
UPLOAD_DIR = Path.home() / ".jarvis" / "uploads"


@dataclass
class DocumentContext:
    content: str
    source: str
    type: str
    pages: int = 1
    tables: int = 0
    rows: int = 0
    chunks: int = 1
    truncated: bool = False

    def to_dict(self) -> dict:
        return {
            "content": self.content,
            "source": self.source,
            "type": self.type,
            "pages": self.pages,
            "tables": self.tables,
            "rows": self.rows,
            "chunks": self.chunks,
            "truncated": self.truncated,
        }


class DocumentProcessor:

    def __init__(self, brain: Any = None, upload_dir: str | Path | None = None):
        self._brain = brain
        self.upload_dir = Path(upload_dir or UPLOAD_DIR)
        self.upload_dir.mkdir(parents=True, exist_ok=True)

    async def save_upload(self, file_bytes: bytes, filename: str) -> str:
        safe = Path(filename).name
        dest = self.upload_dir / f"{datetime.utcnow().strftime('%Y%m%d_%H%M%S')}_{safe}"
        dest.write_bytes(file_bytes)
        return str(dest)

    async def cleanup_old(self, hours: int = 24):
        cutoff = datetime.utcnow() - timedelta(hours=hours)
        for f in self.upload_dir.iterdir():
            try:
                mtime = datetime.utcfromtimestamp(f.stat().st_mtime)
                if mtime < cutoff:
                    f.unlink()
            except Exception as e:
                logger.exception("[DocProc] cleaning up old upload files: %s", e)

    async def process(self, file_path: str) -> DocumentContext:
        ext = Path(file_path).suffix.lower()
        handlers = {
            ".pdf": self._process_pdf,
            ".docx": self._process_docx,
            ".doc": self._process_docx,
            ".xlsx": self._process_xlsx,
            ".xls": self._process_xlsx,
            ".csv": self._process_csv,
            ".txt": self._process_text,
            ".md": self._process_text,
            ".py": self._process_text,
            ".json": self._process_text,
            ".png": self._process_image,
            ".jpg": self._process_image,
            ".jpeg": self._process_image,
            ".webp": self._process_image,
            ".gif": self._process_image,
        }
        handler = handlers.get(ext, self._process_text)
        try:
            ctx = await handler(file_path)
        except Exception as e:
            logger.warning("DocumentProcessor: %s failed: %s — fallback to raw text", ext, e)
            ctx = await self._process_text(file_path)

        if self._estimate_tokens(ctx.content) > TOKEN_LIMIT:
            ctx.content, ctx.chunks = await self._smart_chunk(ctx.content)
            ctx.truncated = True

        return ctx

    async def _process_pdf(self, path: str) -> DocumentContext:
        return self._read_pdf(path)

    @staticmethod
    def _read_pdf(path: str) -> DocumentContext:
        try:
            import pdfplumber
        except ImportError:
            raise ImportError("pdfplumber not installed. Run: pip install pdfplumber")
        pages_text = []
        total_tables = 0
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                text = page.extract_text() or ""
                raw_tables = page.extract_tables() or []
                total_tables += len(raw_tables)
                for tbl in raw_tables:
                    text += "\n\n" + DocumentProcessor._table_to_md(tbl)
                if text.strip():
                    pages_text.append(text.strip())
        return DocumentContext(
            content="\n\n---\n\n".join(pages_text),
            source=path,
            type="pdf",
            pages=len(pages_text),
            tables=total_tables,
        )

    async def _process_docx(self, path: str) -> DocumentContext:
        return self._read_docx(path)

    @staticmethod
    def _read_docx(path: str) -> DocumentContext:
        try:
            from docx import Document
        except ImportError:
            raise ImportError("python-docx not installed. Run: pip install python-docx")
        doc = Document(path)
        paras = [p.text for p in doc.paragraphs if p.text.strip()]
        tables = []
        for tbl in doc.tables:
            rows = [[cell.text for cell in row.cells] for row in tbl.rows]
            tables.append(DocumentProcessor._table_to_md(rows))
        content = "\n".join(paras)
        if tables:
            content += "\n\n## Tables\n\n" + "\n\n".join(tables)
        return DocumentContext(
            content=content,
            source=path,
            type="docx",
            pages=len(paras),
            tables=len(tables),
        )

    async def _process_xlsx(self, path: str) -> DocumentContext:
        return self._read_xlsx(path)

    @staticmethod
    def _read_xlsx(path: str) -> DocumentContext:
        try:
            import pandas as pd
        except ImportError:
            raise ImportError("pandas not installed. Run: pip install pandas openpyxl tabulate")
        xls = pd.ExcelFile(path)
        sheets = []
        total_rows = 0
        for sheet_name in xls.sheet_names:
            df = xls.parse(sheet_name)
            total_rows += len(df)
            try:
                sheet_md = df.head(50).to_markdown(index=False)
            except Exception as e:
                logger.exception("[DocProc] converting xlsx sheet to markdown: %s", e)
                sheet_md = df.head(50).to_string(index=False)
            sheets.append(f"## Sheet: {sheet_name} ({len(df)} rows)\n\n{sheet_md}")
        return DocumentContext(
            content="\n\n".join(sheets),
            source=path,
            type="xlsx",
            tables=len(xls.sheet_names),
            rows=total_rows,
        )

    async def _process_csv(self, path: str) -> DocumentContext:
        return self._read_csv(path)

    @staticmethod
    def _read_csv(path: str) -> DocumentContext:
        try:
            import pandas as pd
        except ImportError:
            raise ImportError("pandas not installed. Run: pip install pandas tabulate")
        df = pd.read_csv(path)
        try:
            preview = df.head(50).to_markdown(index=False)
        except Exception as e:
            logger.exception("[DocProc] converting csv preview to markdown: %s", e)
            preview = df.head(50).to_string(index=False)
        try:
            stats = df.describe().to_markdown()
        except Exception as e:
            logger.exception("[DocProc] generating csv statistics markdown: %s", e)
            stats = df.describe().to_string()
        content = (
            f"## CSV Data — {Path(path).name}\n"
            f"Total rows: {len(df)} | Columns: {list(df.columns)}\n\n"
            f"### Preview (first 50 rows)\n{preview}\n\n"
            f"### Statistics\n{stats}"
        )
        return DocumentContext(
            content=content,
            source=path,
            type="csv",
            rows=len(df),
        )

    async def _process_text(self, path: str) -> DocumentContext:
        return self._read_text(path)

    @staticmethod
    def _read_text(path: str) -> DocumentContext:
        try:
            content = Path(path).read_text(encoding="utf-8", errors="replace")
        except Exception as e:
            content = f"[Could not read file: {e}]"
        ext = Path(path).suffix.lstrip(".")
        return DocumentContext(content=content, source=path, type=ext or "text")

    async def _process_image(self, path: str) -> DocumentContext:
        try:
            b64 = self._read_image_b64(path)
            import importlib as _il
            _llm_router = _il.import_module("core.llm_router")
            vision_result = await _llm_router.complete_vision([{
                "role": "user",
                "content": (
                    "Describe this image in detail. "
                    "Extract ALL visible text (OCR). "
                    "Identify: charts, tables, diagrams, screenshots, photos. "
                    "For charts/graphs: describe the data trend shown. "
                    "For screenshots: describe the UI and any important text."
                ),
                "images": [b64],
            }], timeout=60)
            description = vision_result.unwrap_or("")
        except Exception as e:
            logger.warning("Image processing failed for %s: %s", path, e)
            description = f"[Image: {Path(path).name} — could not process: {e}]"
        return DocumentContext(content=description, source=path, type="image")

    @staticmethod
    def _read_image_b64(path: str) -> str:
        with open(path, "rb") as f:
            return base64.b64encode(f.read()).decode()

    async def _smart_chunk(self, text: str) -> tuple[str, int]:
        chunks = [
            text[i:i + CHUNK_CHARS]
            for i in range(0, len(text), CHUNK_CHARS)
        ][:MAX_CHUNKS]
        if not chunks:
            return text, 1
        logger.info("DocumentProcessor: smart-chunking %d chunks", len(chunks))

        if self._brain is None or not hasattr(self._brain, "reason"):
            try:
                from brain.UnifiedBrain import unified_brain
                self._brain = unified_brain
            except Exception as e:
                logger.exception("[DocProc] loading lazy singleton brain: %s", e)
        if self._brain is None or not hasattr(self._brain, "reason"):
            combined = "\n\n".join(
                f"[Section {i+1}] {c[:500]}..."
                for i, c in enumerate(chunks)
            )
            return combined, len(chunks)

        tasks = [
            self._brain.reason(
                "Summarize this section of a document in 2-3 sentences. "
                "Preserve key facts, numbers, names, and decisions:\n\n" + chunk,
                {},
            )
            for chunk in chunks
        ]
        results = await asyncio.gather(*tasks, return_exceptions=True)

        summaries = []
        for i, r in enumerate(results):
            if isinstance(r, Exception):
                summaries.append(f"[Section {i+1}: summarization failed]")
            else:
                summaries.append(f"[Section {i+1}] {r.answer}")
        return "\n\n".join(summaries), len(chunks)

    @staticmethod
    def _table_to_md(table: list) -> str:
        if not table or not table[0]:
            return ""
        try:
            header = "| " + " | ".join(str(c or "").strip() for c in table[0]) + " |"
            sep = "| " + " | ".join("---" for _ in table[0]) + " |"
            rows = [
                "| " + " | ".join(str(c or "").strip() for c in row) + " |"
                for row in table[1:] if row
            ]
            return "\n".join([header, sep] + rows)
        except Exception as e:
            logger.exception("[DocProc] formatting table to markdown: %s", e)
            return str(table)

    @staticmethod
    def _estimate_tokens(text: str) -> int:
        return len(text) // 4


# Lazy singleton — brain is imported only on first use
_doc_processor_instance: DocumentProcessor | None = None
_doc_processor_lock = threading.Lock()


def get_doc_processor() -> DocumentProcessor:
    global _doc_processor_instance
    if _doc_processor_instance is None:
        with _doc_processor_lock:
            if _doc_processor_instance is None:
                _doc_processor_instance = DocumentProcessor()
    return _doc_processor_instance


class _LazyDocProcessor:
    """Proxy that never imports brain at module level. Brain is only loaded in _smart_chunk."""
    def __getattr__(self, name):
        return getattr(get_doc_processor(), name)


doc_processor = _LazyDocProcessor()
